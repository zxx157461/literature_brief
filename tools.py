"""
tools.py - 文献简报工具集
每个步骤独立封装，供 AI Agent 调度调用

工具列表：
  search_papers          - 搜索 arXiv 论文
  score_papers           - AI 评分筛选论文
  download_paper         - 下载论文 PDF
  extract_images         - 从 PDF 提取所有图片
  extract_text           - 提取 PDF 全文
  generate_text          - 生成文字解读
  generate_caption       - 生成单张图注
  read_paper             - AI 深度阅读理解
  save_to_doc            - 保存到 Word
  generate_poster_script - 文本大模型提炼海报脚本
  generate_poster_image  - 图片大模型生成海报
"""
import json
import base64
import urllib.request
import urllib.error
import urllib.parse
from pathlib import Path
from datetime import datetime
from papers import Paper
from config import cfg, PAPERS_DIR, REPORTS_DIR, SSL_CTX


# ============================================================
# 辅助
# ============================================================

class ToolError(Exception):
    """工具执行失败"""


# ============================================================
# 工具 0：搜索论文
# ============================================================

def search_papers(
    keywords: str = "",
    author: str = "",
    categories: list = None,
    year_from: int = None,
    year_to: int = None,
    max_results: int = 20,
    sort_by: str = "relevance",
) -> dict:
    """
    搜索 arXiv 论文
    输入：搜索参数
    返回：{papers: [Paper对象列表], total: N}
    """
    from searcher import ArxivSearcher, SearchParams

    if not keywords and not author and not categories:
        raise ToolError("至少需要一个搜索条件（关键词/作者/分类）")

    params = SearchParams(
        keywords=keywords,
        author=author,
        categories=categories or [],
        year_from=year_from,
        year_to=year_to,
        max_results=max_results,
        sort_by=sort_by,
    )

    try:
        searcher = ArxivSearcher()
        result = searcher.search(params)
        return {
            "success": True,
            "papers": result.papers,
            "total": result.total_found,
            "query": result.query,
        }
    except Exception as e:
        raise ToolError(f"搜索失败: {e}")


def _load_llm():
    """加载 LLM Caller"""
    from reporter import LLMCaller
    brief_cfg = cfg.get("brief") or {}
    fallback = {
        "provider": "custom",
        "api_key": cfg.get("api_key", ""),
        "base_url": cfg.get("base_url", "https://api.deepseek.com"),
        "model": cfg.get("model", "deepseek-chat"),
        "max_tokens": 2000,
        "temperature": 0.5,
    }
    llm_cfg = brief_cfg if brief_cfg else fallback
    return LLMCaller(llm_cfg)


# ============================================================
# 工具 0.5：AI 评分筛选论文
# ============================================================

def score_papers(
    papers: list,
    topic: str,
    sift_requirement: str = "",
    threshold: float = None,
) -> dict:
    """
    用 AI 对论文列表评分并按阈值筛选
    输入：papers 列表、研究主题、可选筛选要求、可选阈值
    返回：{"filtered": [通过的Paper列表], "scores": [{arxiv_id, score, reason}], "total": N, "passed": M}
    """
    from scorer import PaperScorer

    if threshold is None:
        threshold = cfg.get("auto_score_threshold", 0.6)

    scorer = PaperScorer(topic=topic, sift_requirement=sift_requirement)
    scored = scorer.score_batch(papers)
    filtered = [p for p in scored if p.relevance_score >= threshold]

    return {
        "success": True,
        "filtered": filtered,
        "scores": [
            {"arxiv_id": p.arxiv_id, "score": p.relevance_score, "reason": getattr(p, "score_reason", "")}
            for p in scored
        ],
        "total": len(scored),
        "passed": len(filtered),
        "threshold": threshold,
    }


# ============================================================
# 工具 1：下载论文 PDF
# ============================================================

def download_paper(paper: Paper, force: bool = False) -> dict:
    """
    下载论文 PDF 到本地
    输入：paper 对象（需要 pdf_url 字段）
    返回：{"success": True, "path": "...", "size_kb": 123.4}
    """
    papers_dir = Path(PAPERS_DIR)
    papers_dir.mkdir(parents=True, exist_ok=True)
    save_path = papers_dir / f"{paper.arxiv_id}.pdf"

    if not force and save_path.exists() and save_path.stat().st_size > 1000:
        return {
            "success": True,
            "path": str(save_path),
            "size_kb": save_path.stat().st_size / 1024,
            "cached": True,
        }

    url = paper.pdf_url or f"https://arxiv.org/pdf/{paper.arxiv_id}.pdf"
    headers = {"User-Agent": "Mozilla/5.0 (compatible; literature-brief/1.0)"}
    req = urllib.request.Request(url, headers=headers)

    try:
        with urllib.request.urlopen(req, timeout=60, context=SSL_CTX) as resp:
            content = resp.read()
    except Exception as e:
        raise ToolError(f"下载失败 [{paper.arxiv_id}]: {e}")

    with open(save_path, "wb") as f:
        f.write(content)

    return {
        "success": True,
        "path": str(save_path),
        "size_kb": save_path.stat().st_size / 1024,
        "cached": False,
    }


# ============================================================
# 工具 2：提取 PDF 图片
# ============================================================

def extract_images(paper: Paper, pdf_path: str = None) -> dict:
    """
    从 PDF 提取所有图片，返回元信息列表
    输入：paper 对象，可选 pdf_path（优先用此路径）
    返回：{"success": True, "figures": [FigureInfo列表], "n_total": N}
    """
    import pdfplumber
    from PIL import Image
    import io

    if pdf_path is None:
        pdf_path = str(Path(PAPERS_DIR) / f"{paper.arxiv_id}.pdf")
    path = Path(pdf_path)
    if not path.exists():
        raise ToolError(f"PDF 不存在: {pdf_path}")

    figures_dir = path.parent.parent / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)

    figures = []
    figure_count = 0

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            images = page.images
            if not images:
                continue

            for img_info in images:
                try:
                    x0 = int(img_info["x0"])
                    y0 = int(img_info["top"])
                    x1 = int(img_info["x1"])
                    y1 = int(img_info["bottom"])
                    w, h = x1 - x0, y1 - y0
                    if w < 100 or h < 100:
                        continue

                    pil_img = page.crop((x0, y0, x1, y1), strict=False).to_image()
                    buf = io.BytesIO()
                    pil_img.original.convert("RGB").save(buf, format="JPEG", quality=85)
                    img_jpeg = buf.getvalue()
                    size_kb = len(img_jpeg) / 1024

                    if size_kb < 10:
                        continue

                    figure_count += 1
                    filename = f"{paper.arxiv_id}_p{page_num}_fig{figure_count}.jpg"
                    save_path = figures_dir / filename
                    with open(save_path, "wb") as f:
                        f.write(img_jpeg)

                    figures.append({
                        "figure_id": f"{paper.arxiv_id}_p{page_num}_fig{figure_count}",
                        "page_num": page_num,
                        "path": str(save_path),
                        "width": w,
                        "height": h,
                        "size_kb": round(size_kb, 1),
                    })
                except Exception:
                    continue

    return {"success": True, "figures": figures, "n_total": len(figures)}


# ============================================================
# 工具 3：提取 PDF 全文
# ============================================================

def extract_text(paper: Paper, pdf_path: str = None, max_chars: int = 15000) -> dict:
    """
    从 PDF 提取全文
    输入：paper 对象，可选 pdf_path
    返回：{"success": True, "text": "...", "n_chars": N}
    """
    import re

    if pdf_path is None:
        pdf_path = str(Path(PAPERS_DIR) / f"{paper.arxiv_id}.pdf")
    path = Path(pdf_path)
    if not path.exists():
        raise ToolError(f"PDF 不存在: {pdf_path}")

    import pdfplumber
    texts = []
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            t = page.extract_text() or ""
            t = re.sub(r"\n{3,}", "\n\n", t)
            t = re.sub(r" {2,}", " ", t)
            if t.strip():
                texts.append(f"[第{page_num}页]\n{t.strip()}")

    full_text = "\n\n".join(texts)
    if len(full_text) > max_chars:
        # 截断：保留前中后三段
        chunk = max_chars // 3
        full_text = full_text[:chunk] + "\n\n[...] 中间部分省略 [...]\n\n" + full_text[-chunk:]

    return {
        "success": True,
        "text": full_text,
        "n_chars": len(full_text),
        "truncated": len(full_text) > max_chars,
    }


# ============================================================
# 工具 4：生成文字解读
# ============================================================

REPORT_PROMPT_TEMPLATE = """你是一位专业学术编辑，擅长将论文转化为清晰、深入、有洞见的科研导读。

格式要求：
- 不使用 Markdown 标记（不用 ##、**、-、列表等）
- 用 "序号 | 主题标题" 作为小标题
- 小标题下方紧跟一句"一句话总结"（20字以内，说明这篇论文做了什么）
- 正文用自然段落，首行不缩进，段间空一行
- 每段 3-5 句话
- 结尾列出 3 条研究趋势（用数字序号）

示例格式（模仿学术导读风格）：
===== 示例开始 =====
01 | 论文标题
一句话：XX大学XX等人发现/揭示了XX，可用于XX。

近日，XX大学 XX 等人在 Science 发表研究，发现了XX现象，揭示了XX机制。该工作为理解XX提供了新的理论框架。

总体来看，当前研究的核心趋势在于：
1. ...
2. ...
3. ...
===== 示例结束 =====

论文标题：{title}
作者：{authors}
arXiv：{arxiv_id}

以下是该论文全文或摘要：
{full_text}

请严格按照示例格式，为这篇论文撰写科研导读。"""

def generate_text(paper: Paper, full_text: str) -> dict:
    """
    调用 LLM 生成论文的文字解读
    输入：paper 对象 + 全文（或摘要）
    返回：{"success": True, "text": "...", "n_chars": N}
    """
    llm = _load_llm()

    prompt = REPORT_PROMPT_TEMPLATE.format(
        title=paper.title,
        authors=paper.display_authors,
        arxiv_id=paper.arxiv_id,
        full_text=full_text[:10000],
    )

    try:
        text = llm.chat([
            {"role": "system", "content": "你是一位专业学术编辑，始终严格按格式输出。"},
            {"role": "user", "content": prompt},
        ])
    except Exception as e:
        raise ToolError(f"文字解读生成失败: {e}")

    return {"success": True, "text": text.strip(), "n_chars": len(text)}


# ============================================================
# 工具 5：生成图注
# ============================================================

def generate_caption(
    paper: Paper,
    figure_id: str,
    figure_path: str,
    context: str = None,
) -> dict:
    """
    为一张图生成图注（调用多模态模型看图说话）
    输入：paper 对象 + figure_id + 图片路径 + 可选上下文
    返回：{"success": True, "caption": "...", "figure_id": "..."}
    """
    llm = _load_llm()
    context_text = context or paper.abstract[:500]

    prompt = f"""这是一篇学术论文中的图，请仔细观察并撰写图注。

要求：
1. 1-3 句话
2. 描述图中最关键的发现或数据趋势
3. 指出图中的物理量或实验条件
4. 语言简洁专业

论文标题：{paper.title}
论文背景：{context_text}

请直接输出图注，不要有其他内容。"""

    try:
        # 尝试多模态调用（需要模型支持）
        caption = llm.chat_with_images(prompt, [figure_path])
    except (TypeError, AttributeError):
        # 模型不支持多模态，回退到纯文字
        fallback_prompt = f"""{prompt}

（注意：该模型不支持看图，图注基于上下文推断）
请根据上述论文背景，推断这张图可能的图注内容，描述应尽可能通用。"""
        caption = llm.chat([{"role": "user", "content": fallback_prompt}])

    return {"success": True, "caption": caption.strip(), "figure_id": figure_id}


# ============================================================
# 工具 6：深度阅读理解
# ============================================================

def read_paper(
    paper: Paper,
    full_text: str = "",
) -> dict:
    """
    让 AI 深度阅读理解一篇论文，输出结构化阅读笔记。
    AI 会：理解核心内容 → 提炼关键发现 → 判断是否值得收录 → 决定重点图
    """
    from reporter import LLMCaller

    brief_cfg = cfg.get("brief") or {}
    fallback = {
        "provider": "custom", "api_key": "",
        "base_url": "https://api.deepseek.com",
        "model": "deepseek-chat",
        "max_tokens": 2000, "temperature": 0.3,
    }
    llm = LLMCaller(brief_cfg if brief_cfg else fallback)

    text_content = full_text[:8000] if full_text else paper.abstract

    prompt = f"""你是一位认真阅读论文的研究者。请深度阅读以下论文，输出结构化阅读笔记。

论文：{paper.title}
作者：{paper.display_authors}
arXiv：{paper.arxiv_id}

论文内容：
{text_content}

请严格以 JSON 格式返回（不要有其他内容）：
{{
  "summary": "200字以内的核心内容摘要",
  "key_findings": ["发现1", "发现2", "发现3"],
  "recommended_figures": [],
  "figure_selection_reasons": {{}},
  "is_worth_including": true,
  "reason_for_exclusion": ""
}}"""

    try:
        response = llm.chat([
            {"role": "system", "content": "你是一个严谨的学术阅读助手，始终以纯 JSON 格式回答。"},
            {"role": "user", "content": prompt},
        ])
    except Exception as e:
        raise ToolError(f"阅读理解失败: {e}")

    response = response.strip()
    if response.startswith("```"):
        lines = response.split("\n")
        response = "\n".join(lines[1:] if lines[0].strip() == "```" else lines).rstrip("`")

    try:
        result = json.loads(response)
    except json.JSONDecodeError:
        raise ToolError(f"阅读笔记 JSON 解析失败: {response[:100]}")

    return {"success": True, **result}


# ============================================================
# 工具 7：保存到 Word
# ============================================================

def save_to_doc(
    paper: Paper,
    text: str,
    figures: list = None,
    captions: dict = None,
    output_path: str = None,
) -> dict:
    """
    将文字 + 图片 + 图注保存到 Word 文档
    输入：
      paper 对象
      text: 文字解读
      figures: FigureInfo 列表（可选）
      captions: {figure_id: caption} 字典（可选）
      output_path: 输出路径（可选）
    返回：{"success": True, "path": "..."}
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if figures is None:
        figures = []
    if captions is None:
        captions = {}

    if output_path is None:
        safe_title = "".join(
            c for c in paper.title[:20] if c.isalnum() or c in " -_"
        ).strip()
        output_path = str(
            REPORTS_DIR / f"{paper.arxiv_id}_{safe_title}.docx"
        )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 标题
    doc.add_heading(paper.title, 0)

    # 元信息
    meta = doc.add_paragraph()
    meta.add_run(f"arXiv: {paper.arxiv_id}\n")
    meta.add_run(f"作者: {paper.display_authors}\n")
    meta.add_run(f"日期: {paper.submitted_date}")

    # 正文
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        para = para.replace("**", "").replace("##", "")
        if para:
            doc.add_paragraph(para)

    # 图片 + 图注
    for fig in figures:
        fig_id = fig.get("figure_id", fig.get("id", ""))
        try:
            doc.add_picture(fig["path"], width=doc.sections[0].page_width * 0.65)
            cap_text = captions.get(fig_id, f"图: {fig_id}")
            cap_para = doc.add_paragraph()
            run = cap_para.add_run(cap_text)
            run.italic = True
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"# Warning: 插入图片失败 {fig_id}: {e}")

    doc.save(str(output_path))
    return {"success": True, "path": str(output_path)}


# ============================================================
# 工具 8：提炼海报脚本
# ============================================================

def generate_poster_script(
    paper: Paper,
    full_text: str = "",
    figures: list = None,
) -> dict:
    """
    文本大模型提炼论文核心内容为海报脚本。
    输入：paper 对象 + 全文 + 图片列表
    返回：{script: {poster_title, subtitle, core_findings, visual_description, ...}}
    """
    from poster import PosterTextRefiner, POSTER_TEXT_SYSTEM_PROMPT
    from reporter import LLMCaller
    from config import cfg

    brief_cfg = cfg.get("brief") or {}
    fallback = {
        "provider": "custom",
        "api_key": cfg.get("api_key", ""),
        "base_url": cfg.get("base_url", "https://api.deepseek.com"),
        "model": cfg.get("model", "deepseek-chat"),
        "max_tokens": 2000,
        "temperature": 0.5,
    }
    llm = LLMCaller(brief_cfg if brief_cfg else fallback)

    figures = figures or []
    figures_info = ""
    if figures:
        lines = []
        for i, f in enumerate(figures[:6]):
            if isinstance(f, dict):
                lines.append(
                    f"- 图{i}: {f.get('figure_id', '')}, 第{f.get('page_num', '')}页, "
                    f"尺寸{f.get('width', 0)}x{f.get('height', 0)}px"
                )
            else:
                lines.append(f"- 图{i}: {f.figure_id}, 第{f.page_num}页")
        figures_info = "\n".join(lines)
    else:
        figures_info = "无提取到的图片"

    text_content = full_text[:4000] if full_text else paper.abstract[:4000]

    user_prompt = f"""请为以下论文设计学术海报方案。

论文标题：{paper.title}
作者：{paper.display_authors}
摘要：{paper.abstract[:800]}

论文核心内容（部分）：
{text_content}

论文中的图片：
{figures_info}

请按系统提示词要求的 JSON 格式返回海报脚本。"""

    messages = [
        {"role": "system", "content": POSTER_TEXT_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    response = llm.chat(messages, temperature=0.4)
    response = response.strip()
    if response.startswith("```"):
        lines = response.splitlines()
        response = "\n".join(lines[1:] if lines[0].strip().startswith("```") else lines)
        response = response.rstrip("`").strip()

    try:
        script = json.loads(response)
    except (json.JSONDecodeError, ValueError):
        print(f"# generate_poster_script JSON 解析失败: {response[:200]}")
        script = {
            "poster_title": paper.title[:30],
            "subtitle": "",
            "core_findings": [paper.abstract[:100]],
            "visual_description": "Academic research poster showing key findings and data visualizations from the paper",
            "figure_hints": [],
            "figure_selections": [],
        }

    return {"success": True, "script": script}


# ============================================================
# 工具 9：生成海报图片
# ============================================================

def generate_poster_image(
    poster_script: dict = None,
    reference_image_paths: list = None,
) -> dict:
    """
    图片大模型根据海报脚本和参考图片生成竖版学术海报。
    输入：poster_script dict（含 poster_title/core_findings/visual_description 等）+ reference_image_paths
    返回：{success, arxiv_id, title, path, error?}
    """
    from poster import PosterImageGenerator

    if poster_script is None:
        poster_script = {}

    script = dict(poster_script)
    script["_paper"] = None
    script["_reference_images"] = reference_image_paths or []

    image_gen = PosterImageGenerator()
    return image_gen.generate(script)


# ============================================================
# 工具注册表（供 Agent 查看可用工具）
# ============================================================

TOOL_REGISTRY = {
    "search_papers": {
        "name": "search_papers",
        "description": "搜索 arXiv 论文，返回论文列表。可按关键词/作者/分类/年份搜索。",
        "input": "keywords/author/categories/year_from/year_to/max_results/sort_by",
        "returns": "{papers: [Paper对象列表], total}",
    },
    "score_papers": {
        "name": "score_papers",
        "description": "用 AI 筛选模型对论文列表评分，按相关性阈值过滤。当需要从大量论文中筛选出最相关的论文时调用此工具。",
        "input": "papers 列表, topic 研究主题, sift_requirement 筛选要求（可选）, threshold 阈值（可选，默认读配置）",
        "returns": "{filtered: [通过的Paper列表], scores: [{arxiv_id, score, reason}], total, passed, threshold}",
    },
    "download_paper": {
        "name": "download_paper",
        "description": "下载论文 PDF 到本地。如果 PDF 已存在则跳过（除非 force=True）。",
        "input": "paper 对象（需要有 pdf_url 字段）",
        "returns": "{path, size_kb, cached}",
    },
    "extract_images": {
        "name": "extract_images",
        "description": "从 PDF 提取所有图片，返回图片元信息列表（路径/尺寸/页码）。",
        "input": "paper 对象，pdf_path 可选",
        "returns": "{figures: [{figure_id, page_num, path, width, height, size_kb}], n_total}",
    },
    "extract_text": {
        "name": "extract_text",
        "description": "提取 PDF 全文（自动截断到 max_chars）。",
        "input": "paper 对象，可选 pdf_path 和 max_chars",
        "returns": "{text, n_chars, truncated}",
    },
    "generate_text": {
        "name": "generate_text",
        "description": "调用 LLM 根据全文（或摘要）生成科研导读文字。",
        "input": "paper 对象 + 全文/摘要文本",
        "returns": "{text, n_chars}",
    },
    "generate_caption": {
        "name": "generate_caption",
        "description": "为单张图生成图注（多模态模型看图说话，或回退到上下文推断）。",
        "input": "paper 对象 + figure_id + 图片路径 + 可选上下文",
        "returns": "{caption, figure_id}",
    },
    "read_paper": {
        "name": "read_paper",
        "description": "让 AI 深度阅读理解论文，输出结构化笔记（摘要/关键发现/判断是否收录/推荐重点图）。",
        "input": "paper 对象 + 全文（来自 extract_text）",
        "returns": "{summary, key_findings, recommended_figures, is_worth_including}",
    },
    "save_to_doc": {
        "name": "save_to_doc",
        "description": "将文字+图片+图注保存为 Word 文档。",
        "input": "paper + text + figures列表 + captions字典",
        "returns": "{path}",
    },
    "generate_poster_script": {
        "name": "generate_poster_script",
        "description": "文本大模型提炼论文核心内容为海报脚本。提取关键发现、设计视觉描述、选择参考图。",
        "input": "paper对象 + 全文 + 图片列表",
        "returns": "{script: {poster_title, subtitle, core_findings, visual_description, style_guide, key_visual_elements, figure_selections}}",
    },
    "generate_poster_image": {
        "name": "generate_poster_image",
        "description": "图片大模型根据海报脚本和参考图片生成竖版学术海报。接收系统提示词（规范图片格式）+ 精炼文字 + 参考图片，返回海报图片。",
        "input": "poster_script dict + reference_image_paths list",
        "returns": "{success, arxiv_id, title, path}",
    },
}


def get_tool_schemas() -> list:
    """返回 Agent 可用的工具 schema（OpenAI function calling 格式）"""
    schemas = []
    for tool_name, info in TOOL_REGISTRY.items():
        schemas.append({
            "type": "function",
            "function": {
                "name": tool_name,
                "description": info["description"],
                "parameters": {
                    "type": "object",
                    "properties": {},
                    "required": [],
                },
            },
        })
    return schemas
