"""
reporter.py - 简报生成模块
S11 - 生成图文并茂的 Word 文献简报

完整流程：
  1. 提取 PDF 全文 + 所有图片
  2. AI 筛选重点图 + 生成图注
  3. AI 生成全文详细解读
  4. 组装：文字段落 + 原图 + 图注 → Word

调用 brief LLM（从 config 读取）。
"""
import json
import base64
import socket
import http.client
import urllib.parse
from datetime import datetime
from pathlib import Path
from typing import Optional, Callable
from papers import Paper
from config import cfg, REPORTS_DIR, SSL_CTX
from prompts_manager import prompts


# ============================================================
# LLM 调用层
# ============================================================

class LLMCallError(Exception):
    pass


class LLMCaller:
    def __init__(self, llm_cfg: dict):
        self.cfg = llm_cfg
        self.base_url = llm_cfg["base_url"].rstrip("/")
        self.model = llm_cfg["model"]
        self.api_key = llm_cfg["api_key"]
        self.max_tokens = llm_cfg.get("max_tokens", 2000)
        self.temperature = llm_cfg.get("temperature", 0.5)
        self.timeout = llm_cfg.get("timeout", 600)

    def _get_api_path(self) -> str:
        """根据 base_url 判断 API 路径"""
        # 如果 base_url 已经包含 /chat/completions，不再拼接
        if self.base_url.endswith("/chat/completions"):
            return ""
        # 如果 base_url 以 /v1 或 /v1/ 结尾，拼接 chat/completions
        if self.base_url.rstrip("/").endswith("/v1"):
            return "/chat/completions"
        # 千问兼容模式等：base_url 包含 compatible-mode/v1
        if "/v1" in self.base_url:
            return "/chat/completions"
        provider = self.cfg.get("provider", "openai")
        if provider == "minimax":
            return "/v1/chat/completions"
        return "/chat/completions"

    def _do_request(self, payload: dict) -> dict:
        url = f"{self.base_url}{self._get_api_path()}"
        headers = {
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": f"Bearer {self.api_key}",
        }
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Length"] = str(len(body))

        parsed = urllib.parse.urlparse(url)
        host = parsed.hostname
        port = parsed.port or (443 if parsed.scheme == "https" else 80)
        path = parsed.path + (f"?{parsed.query}" if parsed.query else "")

        if parsed.scheme == "https":
            conn = http.client.HTTPSConnection(host, port, timeout=self.timeout, context=SSL_CTX)
        else:
            conn = http.client.HTTPConnection(host, port, timeout=self.timeout)

        try:
            conn.request("POST", path, body=body, headers=headers)
            resp = conn.getresponse()
            # socket 级超时：防止本地模型生成完成后不关闭连接
            if conn.sock:
                conn.sock.settimeout(self.timeout)
            data = resp.read().decode("utf-8", errors="replace")
            if resp.status >= 400:
                err_msg = f"HTTP {resp.status}: {data[:300]}"
                if resp.status == 404:
                    err_msg = f"API 路径不存在 (404)\n请求地址: {url}\n响应: {data[:200]}"
                elif resp.status == 401:
                    err_msg = f"API Key 无效或未授权 (401)\n请求地址: {url}"
                elif resp.status == 429:
                    err_msg = f"请求过于频繁，请等待 (429)\n请求地址: {url}"
                raise LLMCallError(err_msg)
            return json.loads(data)
        except LLMCallError:
            raise
        except Exception as e:
            raise LLMCallError(f"请求失败: {url}\n错误: {e}")
        finally:
            conn.close()

    def chat(self, messages: list[dict], **kwargs) -> str:
        payload = {
            "model": self.model,
            "messages": messages,
            "max_tokens": kwargs.get("max_tokens", self.max_tokens),
            "temperature": kwargs.get("temperature", self.temperature),
        }
        if self.cfg.get("provider") == "minimax" and payload["temperature"] == 0:
            payload["temperature"] = 0.01

        result = self._do_request(payload)
        if "error" in result:
            raise LLMCallError(f"API error: {result['error']}")
        return result["choices"][0]["message"]["content"]

    def chat_with_tools(self, messages: list[dict], tools: list, **kwargs) -> dict:
        """发送带工具调用的消息，返回完整响应（包含 tool_calls）"""
        payload = {
            "model": self.model,
            "messages": messages,
            "tools": tools,
            "max_tokens": kwargs.get("max_tokens", self.max_tokens),
            "temperature": kwargs.get("temperature", self.temperature),
        }
        if self.cfg.get("provider") == "minimax" and payload["temperature"] == 0:
            payload["temperature"] = 0.01

        result = self._do_request(payload)
        if "error" in result:
            raise LLMCallError(f"API error: {result['error']}")
        return result

    def chat_with_images(
        self, text: str, image_paths: list[str], **kwargs
    ) -> str:
        """
        发送图文混合消息，支持 GPT-4V / DeepSeek-VL 等多模态模型
        image_paths: 本地图片路径列表（会被 base64 编码后发送）
        """
        content = [{"type": "text", "text": text}]
        for path in image_paths:
            with open(path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{b64}"
                }
            })

        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": content}],
            "max_tokens": kwargs.get("max_tokens", 2000),
            "temperature": kwargs.get("temperature", 0.3),
        }

        result = self._do_request(payload)
        if "error" in result:
            raise LLMCallError(f"API error: {result['error']}")
        return result["choices"][0]["message"]["content"]


# ============================================================
# 简报生成 Prompt（从 prompts_manager 读取，支持用户自定义）
# ============================================================

def _get_report_system_prompt() -> str:
    return prompts.get("brief_text")

def _get_report_user_template() -> str:
    return prompts.get("brief_user_template")


# ============================================================
# 图注生成 Prompt
# ============================================================

FIGURE_SELECTION_PROMPT = """论文：{title}
作者：{authors}

这篇论文中共有 {n_figures} 张图。以下是需要筛选的图片（顺序编号）：

{figure_list}

请从中选出最多 {max_select} 张最能体现论文核心发现的图，返回 JSON：

{{
  "selected": [
    {{"figure_num": 1, "reason": "选这张的理由"}},
    ...
  ],
  "reasoning": "筛选思路（1-2句话）"
}}

选图标准：优先选能反映核心物理结论的图（统计曲线、相图、实验结果图），
排除流程图、示意图、作者照片、纯文本页面。"""

FIGURE_CAPTION_PROMPT = """以下是一篇学术论文的图，请仔细观察，然后为其撰写图注。

图注要求：
1. 1-3 句话，描述图中最关键的发现或数据趋势
2. 指出图中标注的物理量或实验条件（如有）
3. 语言简洁专业，符合学术写作风格
4. 不使用 Markdown 格式

论文标题：{title}
论文背景：{context}

请直接输出图注，不要有其他内容。"""


# ============================================================
# 简报生成器
# ============================================================

class PaperReporter:
    """完整图文简报生成器"""

    def __init__(self, topic: str):
        self.topic = topic
        brief_cfg = cfg.get("brief") or {}
        fallback = {
            "provider": "custom",
            "api_key": "",
            "base_url": "https://api.deepseek.com",
            "model": "deepseek-chat",
            "max_tokens": 2000,
            "temperature": 0.5,
        }
        llm_cfg = brief_cfg if brief_cfg else fallback
        self.llm = LLMCaller(llm_cfg)

    def generate_full_report(
        self,
        papers: list[Paper],
        reader,
        on_progress: Optional[Callable] = None,
    ) -> dict:
        """
        生成完整图文简报
        返回 {"success": True, "report_path": "...", "papers_generated": N}
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        results = {"success": True, "papers_generated": 0, "errors": []}

        # ── 组装 Word 文档 ──
        doc = Document()

        # 标题页
        title = doc.add_heading(f"文献简报 — {self.topic}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}\n")
        meta.add_run(f"论文数量: {len(papers)} 篇\n")
        meta.add_run(f"搜索主题: {self.topic}")
        doc.add_paragraph()

        for i, paper in enumerate(papers):
            if on_progress:
                on_progress(f"处理论文 {i+1}/{len(papers)}: {paper.title[:30]}...")

            try:
                self._add_paper_to_doc(
                    doc, paper, reader,
                    paper_index=i + 1,
                    on_progress=on_progress,
                )
                results["papers_generated"] += 1
            except Exception as e:
                err_msg = f"[{paper.arxiv_id}] {e}"
                results["errors"].append(err_msg)
                print(f"# Warning: {err_msg}")

        # 保存
        date_str = datetime.now().strftime("%Y-%m-%d")
        safe_topic = "".join(
            c for c in self.topic[:20] if c.isalnum() or c in " -_"
        ).strip()
        output_path = Path(REPORTS_DIR) / f"简报_{safe_topic}_{date_str}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        results["report_path"] = str(output_path)

        return results

    def _add_paper_to_doc(
        self, doc, paper: Paper, reader,
        paper_index: int, on_progress: Optional[callable] = None,
    ):
        """把一篇论文的文字+图片添加到 Word 文档"""
        # ── 下载 PDF ──
        if on_progress:
            on_progress(f"  下载 PDF...")
        pdf_path = Path(reader.papers_dir) / f"{paper.arxiv_id}.pdf"
        if not pdf_path.exists():
            try:
                self._download_pdf(paper, pdf_path)
            except Exception as e:
                print(f"# Warning: PDF 下载失败 [{paper.arxiv_id}]: {e}")

        # ── 提取全文 ──
        if on_progress:
            on_progress(f"  提取全文...")
        try:
            full_text = reader.read_full_text(paper)
        except Exception as e:
            full_text = paper.abstract  # 回退到摘要
            print(f"# Warning: 全文提取失败 [{paper.arxiv_id}]，使用摘要: {e}")

        # ── 生成文字解读 ──
        if on_progress:
            on_progress(f"  生成文字解读...")
        try:
            briefing_text = self._generate_text(paper, full_text)
        except Exception as e:
            briefing_text = f"[文字解读生成失败: {e}]"
            print(f"# Warning: {briefing_text}")

        # ── 提取图片 ──
        figures = []
        if on_progress:
            on_progress(f"  提取图片...")
        try:
            figures = reader.extract_all_figures(paper)
        except Exception as e:
            print(f"# Warning: 图片提取失败 [{paper.arxiv_id}]: {e}")

        # ── AI 筛选重点图 ──
        selected_figures = []
        if figures:
            if on_progress:
                on_progress(f"  AI 筛选重点图...")
            try:
                selected_figures = self._select_figures(paper, figures)
            except Exception as e:
                print(f"# Warning: 图片筛选失败 [{paper.arxiv_id}]: {e}")

        # ── AI 生成图注 ──
        figures_with_captions = []
        for fig in selected_figures:
            if on_progress:
                on_progress(f"  生成图注: {fig.figure_id}...")
            try:
                caption = self._generate_figure_caption(paper, fig, full_text)
                figures_with_captions.append((fig, caption))
            except Exception as e:
                figures_with_captions.append((fig, f"[图注生成失败]"))

        # ── 组装到 Word ──
        # 论文标题
        doc.add_heading(f"{paper_index:02d} | {paper.title[:60]}", level=2)

        # 文字解读（分段）
        for para in briefing_text.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            # 跳过 Markdown 残留
            para = para.replace("**", "").replace("##", "").replace("- ", "")
            if para:
                doc.add_paragraph(para)

        doc.add_paragraph()

        # 图片 + 图注
        for fig, caption in figures_with_captions:
            try:
                from docx.shared import Inches
                # 按原始尺寸插入，限制最大宽度为页面宽度的 90%
                page_width = doc.sections[0].page_width
                max_width = page_width * 0.9
                doc.add_picture(fig.image_path, width=max_width)
                cap_para = doc.add_paragraph()
                run = cap_para.add_run(f"图{fig.figure_id.split('_')[-1]}: {caption}")
                run.italic = True
                cap_para.paragraph_format.alignment = 1  # 居中
                doc.add_paragraph()
            except Exception as e:
                print(f"# Warning: 插入图片失败 {fig.image_path}: {e}")

        doc.add_paragraph("─" * 40)
        doc.add_paragraph()

    def _download_pdf(self, paper: Paper, save_path: Path) -> None:
        """下载论文 PDF"""
        import urllib.request
        from config import SSL_CTX

        url = paper.pdf_url or f"https://arxiv.org/pdf/{paper.arxiv_id}.pdf"
        headers = {"User-Agent": "Mozilla/5.0 (compatible; literature-brief/1.0)"}
        req = urllib.request.Request(url, headers=headers)

        save_path.parent.mkdir(parents=True, exist_ok=True)

        with urllib.request.urlopen(req, timeout=60, context=SSL_CTX) as resp:
            content = resp.read()
        with open(save_path, "wb") as f:
            f.write(content)

    def _generate_text(self, paper: Paper, full_text: str) -> str:
        """调用 LLM 生成文字解读"""
        # 截断过长的全文（避免超出 Token 限制）
        # 本地模型速度较慢，进一步压缩输入长度
        is_local = "localhost" in self.llm.base_url or "127.0.0.1" in self.llm.base_url
        max_input = 3000 if is_local else 6000
        max_text = 4000 if is_local else 8000

        text_to_send = full_text
        if len(text_to_send) > max_text:
            text_to_send = full_text[:max_text // 2] + "\n\n[... 中间部分省略 ...]\n\n" + full_text[-max_text // 2:]

        prompt = _get_report_user_template().format(
            title=paper.title,
            authors=paper.display_authors,
            arxiv_id=paper.arxiv_id,
            submitted_date=paper.submitted_date,
            abstract=paper.abstract[:500] if is_local else paper.abstract[:1000],
            full_text=text_to_send[:max_input],
        )

        messages = [
            {"role": "system", "content": _get_report_system_prompt()},
            {"role": "user", "content": prompt},
        ]

        return self.llm.chat(messages)

    def _select_figures(self, paper: Paper, figures: list) -> list:
        """调用多模态 LLM 看图筛选重点图"""
        if not figures:
            return []

        # 优先尝试多模态看图筛选
        try:
            return self._select_figures_with_vision(paper, figures)
        except Exception as e:
            print(f"# Warning: 多模态看图筛选失败，回退到元信息筛选: {e}")
            return self._select_figures_by_meta(paper, figures)

    def _select_figures_with_vision(self, paper: Paper, figures: list) -> list:
        """用多模态模型真正看图后筛选"""
        max_select = 4

        # 构建带图的 prompt
        content = [
            {"type": "text", "text": f"""论文：{paper.title}
作者：{paper.display_authors}

这篇论文共有 {len(figures)} 张图。请仔细查看每张图，从中选出最多 {max_select} 张最能体现论文核心发现的图。

选图标准（按优先级）：
1. 直接展示核心实验结果或数据趋势的图（如统计曲线、相图、测量结果）
2. 包含关键物理量或定量关系的图
3. 图中有明确标注、坐标轴、数据点的图
4. 排除纯示意图、流程图、系统架构图、作者照片

请严格返回 JSON（不要 markdown）：
{{
  "selected": [
    {{"figure_num": 1, "reason": "这张展示了..."}},
    ...
  ],
  "reasoning": "总体筛选思路"
}}"""}
        ]

        # 把所有图加入消息（base64 编码）
        for i, fig in enumerate(figures, 1):
            with open(fig.image_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            content.append({"type": "text", "text": f"图{i}:"})
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"}
            })

        payload = {
            "model": self.llm.model,
            "messages": [{"role": "user", "content": content}],
            "max_tokens": 2000,
            "temperature": 0.3,
        }

        result = self.llm._do_request(payload)
        response = result["choices"][0]["message"]["content"]

        # 解析 JSON
        response = response.strip()
        if response.startswith("```"):
            lines = response.split("\n")
            response = "\n".join(lines[1:] if lines[0].strip() == "```" else lines).rstrip("`")

        try:
            data = json.loads(response)
        except json.JSONDecodeError:
            print(f"# Warning: 多模态图筛选 JSON 解析失败: {response[:200]}")
            return figures[:3]

        selected = []
        num_to_fig = {i + 1: f for i, f in enumerate(figures)}
        for item in data.get("selected", []):
            fig_num = item.get("figure_num")
            if fig_num in num_to_fig:
                selected.append(num_to_fig[fig_num])

        return selected if selected else figures[:3]

    def _select_figures_by_meta(self, paper: Paper, figures: list) -> list:
        """回退：只用元信息筛选（不看图内容）"""
        figure_list = "\n".join(
            f"图{i+1}: {f.figure_id} (第{f.page_num}页, "
            f"{f.width}×{f.height}px, {f.size_kb:.0f}KB)"
            for i, f in enumerate(figures)
        )

        prompt = FIGURE_SELECTION_PROMPT.format(
            title=paper.title,
            authors=paper.display_authors,
            n_figures=len(figures),
            figure_list=figure_list,
            max_select=4,
        )

        response = self.llm.chat([
            {"role": "system", "content": "你是一个严谨的学术图片筛选助手，始终以 JSON 格式回答。"},
            {"role": "user", "content": prompt},
        ])

        response = response.strip()
        if response.startswith("```"):
            lines = response.split("\n")
            response = "\n".join(lines[1:] if lines[0].strip() == "```" else lines).rstrip("`")

        try:
            data = json.loads(response)
        except json.JSONDecodeError:
            print(f"# Warning: 图筛选 JSON 解析失败，使用前3张: {response[:100]}")
            return figures[:3]

        selected = []
        num_to_fig = {i + 1: f for i, f in enumerate(figures)}
        for item in data.get("selected", []):
            fig_num = item.get("figure_num")
            if fig_num in num_to_fig:
                selected.append(num_to_fig[fig_num])

        return selected if selected else figures[:3]

    def _generate_figure_caption(self, paper: Paper, fig, full_text: str) -> str:
        """调用 LLM 看图并生成图注"""
        prompt = FIGURE_CAPTION_PROMPT.format(
            title=paper.title,
            context=paper.abstract[:500],
        )

        # 发给多模态模型（如果支持）
        try:
            caption = self.llm.chat_with_images(
                prompt, [fig.image_path]
            )
            return caption.strip()
        except (TypeError, AttributeError):
            # 如果 chat_with_images 不存在（模型不支持），回退
            pass

        # 回退：基于全文上下文生成图注
        context_snippet = (
            f"论文探讨了 {paper.title}，"
            f"该图位于第{fig.page_num}页，是论文的核心结果图之一。"
            f"结合论文内容，图注应描述该图所展示的关键数据趋势或物理现象。"
        )
        fallback_prompt = f"""{context_snippet}

图注要求：1-2句话，描述图中最关键的发现。
请直接输出图注，不要有其他内容。"""

        try:
            return self.llm.chat([
                {"role": "user", "content": fallback_prompt}
            ]).strip()
        except LLMCallError:
            return f"第{fig.page_num}页核心结果图"


# ============================================================
# 快速测试（需要先下载 PDF）
# ============================================================
if __name__ == "__main__":
    from reader import PaperReader
    from papers import Paper

    paper = Paper(
        arxiv_id="2301.09550",
        title="Planktonic Active Matter",
        authors=["Anupam Sengupta"],
        abstract="This paper presents a comprehensive study...",
        categories=["cond-mat.soft"],
        submitted_date="2023-01-01",
        updated_date="2023-01-01",
        pdf_url="",
    )

    reader = PaperReader()
    print(f"PDF 目录: {reader.papers_dir}")
    print(f"图片目录: {reader.figures_dir}")

    # 提取全文测试
    try:
        text = reader.read_full_text(paper)
        print(f"全文长度: {len(text)} 字")
        print(f"前200字: {text[:200]}")
    except Exception as e:
        print(f"全文提取: {e}")

    # 提取图片测试
    try:
        figs = reader.extract_all_figures(paper)
        print(f"发现 {len(figs)} 张图片")
        for f in figs[:3]:
            print(f"  {f.figure_id}: {f.size_kb:.0f}KB, {f.width}×{f.height}")
    except Exception as e:
        print(f"图片提取: {e}")
